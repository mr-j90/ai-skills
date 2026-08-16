#!/usr/bin/env python3
"""
Render a canonical walkthrough markdown file to a Word (.docx) document.

Usage:
    python3 render_docx.py INPUT.md OUTPUT.docx [--accent-color #HEX]

Uses the same canonical markdown format as render_pdf.py so both outputs
stay consistent.
"""

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# --- Markdown parsing (mirrors render_pdf.py) ----------------------------

def parse_markdown(md_text: str, base_dir: Path):
    blocks = []
    lines = md_text.splitlines()
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            blocks.append({"type": "title", "text": stripped[2:].strip()})
            j = i + 1
            while j < n and not lines[j].strip():
                j += 1
            if j < n and not lines[j].strip().startswith("#"):
                subtitle = []
                while j < n and lines[j].strip() and not lines[j].strip().startswith("#"):
                    subtitle.append(lines[j].strip())
                    j += 1
                blocks.append({"type": "subtitle", "text": " ".join(subtitle)})
                i = j
                continue
            i += 1
            continue

        if stripped.startswith("## "):
            blocks.append({"type": "h2", "text": stripped[3:].strip()})
            i += 1
            continue

        if stripped.startswith("### "):
            blocks.append({"type": "h3", "text": stripped[4:].strip()})
            i += 1
            continue

        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m:
            alt, path = m.group(1), m.group(2)
            blocks.append({
                "type": "image",
                "alt": alt,
                "path": str((base_dir / path).resolve()),
            })
            i += 1
            continue

        m = re.match(r">\s*\[!(TIP|NOTE|WARNING)\]\s*$", stripped)
        if m:
            kind = m.group(1)
            body = []
            j = i + 1
            while j < n and lines[j].strip().startswith(">"):
                body.append(lines[j].strip().lstrip(">").strip())
                j += 1
            blocks.append({
                "type": "callout",
                "kind": kind,
                "text": " ".join(l for l in body if l),
            })
            i = j
            continue

        if stripped.startswith("- "):
            items = []
            while i < n and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append({"type": "bullets", "items": items})
            continue

        para = [stripped]
        i += 1
        while i < n:
            l = lines[i].strip()
            if not l:
                break
            if l.startswith("#") or l.startswith(">") or l.startswith("- "):
                break
            if re.match(r"!\[([^\]]*)\]\(([^)]+)\)", l):
                break
            para.append(l)
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para)})

    return blocks


# --- Inline formatting ---------------------------------------------------

def add_runs(paragraph, text: str, base_font="Calibri", base_size=11,
             base_color=RGBColor(0x1f, 0x29, 0x37)):
    """Parse **bold** and `code` markers and emit runs. Links become 'label (url)'."""
    # flatten links
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)

    # split on **...** and `...`
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = base_color
        else:
            run = paragraph.add_run(part)
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color


# --- Cell shading helper -------------------------------------------------

def shade_cell(cell, hex_color: str):
    """Set cell background shading (e.g. 'ECFDF5')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def set_cell_border_left(cell, hex_color: str, size_pt=3):
    """Add a thick left border to a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_pt * 8))  # size in eighths of a point
    left.set(qn("w:color"), hex_color.lstrip("#"))
    # remove any existing left child
    existing = tc_borders.find(qn("w:left"))
    if existing is not None:
        tc_borders.remove(existing)
    tc_borders.append(left)


# --- Callout renderer ----------------------------------------------------

CALLOUT_COLORS = {
    "TIP": {"bg": "ECFDF5", "border": "10B981"},
    "NOTE": {"bg": "EFF6FF", "border": "3B82F6"},
    "WARNING": {"bg": "FEF3C7", "border": "F59E0B"},
}


def add_callout(doc: Document, kind: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    cell.width = Inches(6.3)

    colors = CALLOUT_COLORS[kind]
    shade_cell(cell, colors["bg"])
    set_cell_border_left(cell, colors["border"], size_pt=3)

    # label
    label_p = cell.paragraphs[0]
    label_p.paragraph_format.space_before = Pt(4)
    label_p.paragraph_format.space_after = Pt(2)
    label_run = label_p.add_run(kind)
    label_run.bold = True
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)

    # body
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(6)
    add_runs(body_p, text, base_size=10)

    # trailing spacer
    doc.add_paragraph()


# --- Main render ---------------------------------------------------------

def hex_to_rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def render(md_path: Path, docx_path: Path, accent_hex: str):
    md_text = md_path.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text, md_path.parent)

    doc = Document()

    # margins
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    accent = hex_to_rgb(accent_hex)
    ink = RGBColor(0x1f, 0x29, 0x37)
    muted = RGBColor(0x6b, 0x72, 0x80)

    for b in blocks:
        t = b["type"]

        if t == "title":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(b["text"])
            run.bold = True
            run.font.name = "Cambria"
            run.font.size = Pt(24)
            run.font.color.rgb = ink

        elif t == "subtitle":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(14)
            run = p.add_run(b["text"])
            run.italic = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = muted

        elif t == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(b["text"])
            run.bold = True
            run.font.name = "Cambria"
            run.font.size = Pt(16)
            run.font.color.rgb = accent

        elif t == "h3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(b["text"])
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.font.color.rgb = ink

        elif t == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_runs(p, b["text"])

        elif t == "bullets":
            for item in b["items"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                add_runs(p, item)

        elif t == "image":
            path = b["path"]
            if os.path.exists(path):
                try:
                    doc.add_picture(path, width=Inches(6.0))
                    # caption
                    if b["alt"]:
                        cap = doc.add_paragraph()
                        cap.paragraph_format.space_after = Pt(10)
                        run = cap.add_run(b["alt"])
                        run.italic = True
                        run.font.name = "Calibri"
                        run.font.size = Pt(9)
                        run.font.color.rgb = muted
                except Exception as e:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[Image error: {e}]")
                    run.italic = True
                    run.font.color.rgb = muted
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[Missing image: {os.path.basename(path)}]")
                run.italic = True
                run.font.color.rgb = muted

        elif t == "callout":
            add_callout(doc, b["kind"], b["text"])

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main():
    p = argparse.ArgumentParser()
    p.add_argument("input")
    p.add_argument("output")
    p.add_argument("--accent-color", default="#2563eb")
    args = p.parse_args()

    md_path = Path(args.input).resolve()
    docx_path = Path(args.output).resolve()

    if not md_path.exists():
        print(f"Input markdown not found: {md_path}", file=sys.stderr)
        sys.exit(1)

    render(md_path, docx_path, args.accent_color)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main()
